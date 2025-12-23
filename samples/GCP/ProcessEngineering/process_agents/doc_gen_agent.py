# process_agents/doc_gen_agent.py
from google.adk.agents import LlmAgent
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import re
import time
import networkx as nx
import matplotlib.pyplot as plt

def generate_clean_diagram(process_name: str, edge_list_json: str) -> str:
    """Generates a professional process flowchart using the standard template style."""
    try:
        time.sleep(2)
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
        filename = os.path.join(output_dir, f"{process_name.lower().replace(' ', '_')}_flow.png")
        
        # Robustly parse edge list
        clean_edges = re.sub(r'^```json\s*|```$', '', edge_list_json.strip(), flags=re.MULTILINE)
        edges = json.loads(clean_edges)
        
        G = nx.DiGraph()
        G.add_edges_from(edges)
        
        plt.figure(figsize=(12, 7))
        pos = nx.kamada_kawai_layout(G) 
        
        nx.draw(G, pos, with_labels=True, node_color='#FDFEFE', 
                edge_color='#34495E', node_size=3000, font_size=9, 
                font_weight='bold', arrows=True, arrowsize=20, 
                width=1.5, node_shape='s', edgecolors="#34495E")
        
        plt.title(f"Process Workflow: {process_name}", fontsize=14, pad=20)
        plt.savefig(filename, bbox_inches='tight', dpi=300)
        plt.close()
        return filename
    except Exception as e:
        print(f"Diagram Error: {e}")
        return "None"

def create_standard_doc(process_name: str, content_json: str, diagram_path: str) -> str:
    """Compiles a professional process specification following the hierarchical standard."""
    try:
        print(f"create_standard_doc called with:")
        print(f"================================")
        print(f"  process_name: {process_name}")
        print(f"  content_json: {content_json}")
        print(f"  diagram_path: {diagram_path}")
        print(f"================================")

        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
        doc_path = os.path.join(output_dir, f"{process_name.lower().replace(' ', '_')}_specification.docx")
        
        # Clean and parse content JSON
        clean_content = re.sub(r'^```json\s*|```$', '', content_json.strip(), flags=re.MULTILINE)
        data = json.loads(clean_content)
        
        doc = docx.Document()
        
        # --- TITLE PAGE ---
        doc.add_paragraph("\n" * 3)
        p_title = doc.add_heading(f"{process_name}\nProcess Description Document", 0)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("\n" * 4)
        doc.add_paragraph("Standard Process Description Guideline").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Automated Generation").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # --- FOREWORD ---
        doc.add_heading('FOREWORD', level=1)
        doc.add_paragraph(data.get("Foreword", "This document defines the structured workflow for the identified process."))
        
        # --- 1.0 PURPOSE ---
        doc.add_heading('1.0 PURPOSE', level=1)
        doc.add_paragraph(data.get("Purpose", "To establish clear operational guidelines and objectives."))

        # --- 2.0 APPLICATION ---
        doc.add_heading('2.0 APPLICATION', level=1)
        app_data = data.get("Application", {})
        
        doc.add_heading('2.1 Process Elements', level=2)
        doc.add_paragraph(app_data.get("Elements", "Core activities and decision points."))
        
        doc.add_heading('2.2 Associated Documents', level=2)
        doc.add_paragraph(app_data.get("DocumentTypes", "Required artifacts and output logs."))

        # --- 3.0 PROCESS FLOWCHARTS ---
        doc.add_heading('3.0 PROCESS FLOWCHARTS', level=1)
        if diagram_path != "None" and os.path.exists(diagram_path):
            doc.add_paragraph("3.1 Functional Process Flow")
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(diagram_path, width=Inches(6.0))
        
        # --- APPENDIX A: GLOSSARY ---
        doc.add_page_break()
        doc.add_heading('APPENDIX A: GLOSSARY', level=1)
        glossary = data.get("Glossary", {})
        for term, definition in glossary.items():
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{term}: ").bold = True
            p.add_run(definition)

        doc.save(doc_path)
        return f"SUCCESS: Process document generated at {doc_path}"
    except Exception as e:
        return f"ERROR: Documentation failed - {str(e)}"

doc_gen_agent = LlmAgent(
    name='Documentation_Agent',
    model='gemini-2.0-flash-001',
    description='Generates standardized process specifications.',
    instruction=(
        "You are a Senior Technical Architect. You will receive a JSON object containing a detailed process design. "
        "Your goal is to ensure ZERO DATA LOSS when moving information from JSON to the Word document. "
        "1. CALL 'generate_clean_diagram': Use the 'steps' or 'workflow' array from the JSON to create pairs." \
        "   You MUST call 'generate_clean_diagram' using a list of PAIRS (2-tuples). "
        "   Example: [['Step 1', 'Step 2'], ['Step 2', 'Step 3']]. "
        "   Do NOT pass a single list of strings. This prevents the 'Edge tuple' error. "
        "2. CALL 'create_standard_doc': You MUST map the JSON content to these exact parameters: "
        "   - 'process_name': The title of the process. "
        "   - 'content_json': Construct this by extracting the following from the source JSON: "
        "     - 'Foreword': Use the full 'process_description' or 'background' text. "
        "     - 'Purpose': Include all 'goals', 'objectives', and 'success_metrics'. "
        "     - 'Application': Under 'Elements', list every phase identified. Under 'DocumentTypes', list every artifact. "
        "     - 'Glossary': Include every term defined in the source JSON. "
        "CRITICAL: Do not summarize. If the source JSON has 30 steps, the document must reflect the detail of all 30 steps. "
        "Only output the 'SUCCESS' message from the tool call."
    ),
    tools=[generate_clean_diagram, create_standard_doc]
)