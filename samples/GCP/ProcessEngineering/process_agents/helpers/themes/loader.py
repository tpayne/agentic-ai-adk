# process_agents/helpers/themes/loader.py

import json
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


THEMES_DIR = os.path.dirname(__file__)


# ------------------------------------------------------------
# Load theme JSON safely
# ------------------------------------------------------------
def load_theme(theme_name: str) -> dict:
    """
    Load a theme JSON file from the themes directory.

    Returns a dict containing the theme definition.
    Raises clear errors for missing or invalid files.
    """
    filename = f"{theme_name}.json"
    path = os.path.join(THEMES_DIR, filename)

    if not os.path.exists(path):
        raise FileNotFoundError(f"Theme file not found: {path}")

    with open(path, "r", encoding="utf-8") as f:
        theme = json.load(f)

    # Ignore documentation blocks
    theme.pop("__doc__", None)

    # Validate required keys
    required = ["fonts", "colors", "heading_sizes", "body_size"]
    for key in required:
        if key not in theme or theme[key] is None:
            raise ValueError(f"Theme '{theme_name}' is missing required key: '{key}'")

    return theme


# ------------------------------------------------------------
# Safe color application
# ------------------------------------------------------------
def _apply_color(style, rgb_hex: str | None) -> None:
    """
    Apply a hex RGB color to a style if rgb_hex is valid.
    Null or empty values are ignored safely.
    """
    if not rgb_hex:
        return
    try:
        style.font.color.rgb = RGBColor.from_string(rgb_hex)
    except Exception:
        # Invalid hex string — ignore silently
        pass


# ------------------------------------------------------------
# Apply theme to document
# ------------------------------------------------------------
def apply_theme(doc, theme_name: str = "corporate_standard") -> None:
    """
    Apply a named theme to a python-docx Document.

    This function modifies:
    - Normal style
    - Heading 1–5 styles
    - Table Grid style
    - Footer text (first section only)
    """
    theme = load_theme(theme_name)

    fonts = theme["fonts"]
    colors = theme["colors"]
    heading_sizes = theme["heading_sizes"]
    body_size = theme["body_size"]

    # -------------------------
    # Normal / body text
    # -------------------------
    try:
        normal = doc.styles["Normal"]
        normal.font.name = fonts.get("body", "Calibri")
        normal.font.size = Pt(body_size)
    except KeyError:
        pass

    # -------------------------
    # Headings (H1–H5)
    # -------------------------
    def style_heading(name: str, size_key: str, color_key: str | None = None):
        try:
            style = doc.styles[name]
        except KeyError:
            return

        style.font.name = fonts.get("heading", "Segoe UI")
        style.font.size = Pt(heading_sizes[size_key])
        style.font.bold = True

        if color_key:
            _apply_color(style, colors.get(color_key))

    style_heading("Heading 1", "h1", "primary")
    style_heading("Heading 2", "h2", "primary")
    style_heading("Heading 3", "h3", "secondary")
    style_heading("Heading 4", "h4")
    style_heading("Heading 5", "h5")

    # -------------------------
    # Table style
    # -------------------------
    try:
        table_style = doc.styles["Table Grid"]
        table_style.font.name = fonts.get("body", "Calibri")
        table_style.font.size = Pt(body_size)
    except KeyError:
        pass

    # -------------------------
    # Footer text (first section only)
    # -------------------------
    footer_text = theme.get("footer_text")
    if footer_text:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = footer_text
        p.style = doc.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
