# process_agents/helpers/doc_technical.py

import docx
import json
import os
import traceback
import logging

from docx.shared import Inches, Pt

from .doc_structure import (
    apply_iso_table_formatting,
)
from ..step_diagram_agent import generate_step_diagram_for_step

logger = logging.getLogger("ProcessArchitect.DocTechnical")


# ============================================================
# 5.0 METRICS & KPIs
# ============================================================

def _add_metrics_section(doc: docx.Document, metrics) -> None:
    """
    5.0 Metrics & KPIs — ISO formatted.
    Supports:
      - list of strings
      - list of dicts with keys: name, description, measurement_frequency, target, sub_metrics
    """
    try:
        if not isinstance(metrics, list) or not metrics:
            return

        doc.add_heading("5.0 Metrics and Key Performance Indicators (KPIs)", level=1)
        doc.add_paragraph(
            "The following metrics help monitor performance and ensure alignment with business objectives."
        )

        # Simple list of strings
        if all(isinstance(m, str) for m in metrics):
            for m in metrics:
                doc.add_paragraph(m, style="List Bullet")
            doc.add_paragraph()
            return

        # Table for structured metrics
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Description"
        hdr[2].text = "Measurement / Frequency"
        hdr[3].text = "Target"

        for idx, m in enumerate(metrics, start=1):
            if isinstance(m, str):
                row = table.add_row().cells
                row[0].text = m
                continue

            if not isinstance(m, dict):
                continue

            name = m.get("name") or m.get("metric_name") or f"Metric {idx}"
            description = m.get("description", "")
            measurement = (
                m.get("measurement")
                or m.get("measurement_frequency")
                or ""
            )
            target = m.get("target", "")

            row = table.add_row().cells
            row[0].text = str(name)
            row[1].text = str(description)
            row[2].text = str(measurement)
            row[3].text = str(target)

            # Sub-metrics
            sub_metrics = m.get("sub_metrics", [])
            if isinstance(sub_metrics, list) and sub_metrics:
                sub_row = table.add_row().cells
                sub_row[0].merge(sub_row[1]).merge(sub_row[2]).merge(sub_row[3])
                p = sub_row[0].paragraphs[0]
                run = p.add_run("Sub-metrics:\n")
                run.bold = True

                for sm in sub_metrics:
                    if isinstance(sm, str):
                        p.add_run(f"• {sm}\n")
                    elif isinstance(sm, dict):
                        line = sm.get("metric_name", "Sub-metric")
                        if "description" in sm:
                            line = f"{line} – {sm['description']}"
                        p.add_run(f"• {line}\n")

        apply_iso_table_formatting(table)
        doc.add_paragraph()

    except Exception:
        traceback.print_exc()


# ============================================================
# 9.0 SYSTEM REQUIREMENTS
# ============================================================

def _add_system_requirements(doc: docx.Document, system_requirements) -> None:
    """
    9.0 System Requirements — ISO formatted.
    Supports:
      - list of dicts
      - list of simple values
      - mixed/nested structures
    """
    try:
        if not isinstance(system_requirements, list) or not system_requirements:
            return

        doc.add_heading("9.0 System Requirements", level=1)
        doc.add_paragraph(
            "The following system requirements are essential for the successful implementation of this process."
        )

        # Case 1: List of dictionaries
        if all(isinstance(item, dict) for item in system_requirements):

            all_keys = set()
            for item in system_requirements:
                all_keys.update(item.keys())

            ordered_keys = []
            if "name" in all_keys:
                ordered_keys.append("name")
            if "details" in all_keys:
                ordered_keys.append("details")

            remaining = [k for k in all_keys if k not in ("name", "details")]
            ordered_keys.extend(sorted(remaining))

            table = doc.add_table(rows=1, cols=len(ordered_keys))
            hdr = table.rows[0].cells

            for idx, key in enumerate(ordered_keys):
                hdr[idx].text = key.replace("_", " ").title()

            for item in system_requirements:
                row = table.add_row().cells
                for idx, key in enumerate(ordered_keys):
                    val = item.get(key, "")
                    if isinstance(val, dict):
                        row[idx].text = "; ".join(f"{k}: {v}" for k, v in val.items())
                    elif isinstance(val, list):
                        row[idx].text = ", ".join(str(x) for x in val)
                    else:
                        row[idx].text = str(val)

            apply_iso_table_formatting(table)
            doc.add_paragraph()
            return

        # Case 2: Simple list
        if all(isinstance(item, (str, int, float)) for item in system_requirements):
            for item in system_requirements:
                doc.add_paragraph(str(item), style="List Bullet")
            doc.add_paragraph()
            return

        # Case 3: Mixed or nested
        doc.add_paragraph(
            "The system requirements contain mixed or nested data. "
            "The following section renders them in a readable hierarchical format."
        )

        def render_recursive(value, level=0):
            indent_style = "List Bullet" if level > 0 else "Normal"

            if isinstance(value, dict):
                for k, v in value.items():
                    p = doc.add_paragraph(style=indent_style)
                    r = p.add_run(f"{k.replace('_',' ').title()}: ")
                    r.bold = True
                    if isinstance(v, (dict, list)):
                        render_recursive(v, level + 1)
                    else:
                        p.add_run(str(v))

            elif isinstance(value, list):
                for item in value:
                    if isinstance(item, (dict, list)):
                        render_recursive(item, level + 1)
                    else:
                        doc.add_paragraph(str(item), style=indent_style)

            else:
                doc.add_paragraph(str(value), style=indent_style)

        render_recursive(system_requirements)
        doc.add_paragraph()

    except Exception:
        traceback.print_exc()


# ============================================================
# 10.0 PROCESS FLOW DIAGRAM
# ============================================================

def _add_flowchart_section(doc: docx.Document, process_name: str) -> None:
    """10.0 Flow Diagram — ISO formatted."""
    try:
        diag_file = f"output/{process_name.lower().replace(' ', '_')}_flow.png"
        fallback = "output/process_flow.png"

        if not os.path.exists(diag_file):
            if not os.path.exists(fallback):
                return
            diag_file = fallback

        doc.add_heading("10.0 Process Flow Diagram", level=1)
        doc.add_paragraph(
            "The following diagram provides a high-level visualization of the process flow."
        )
        doc.add_picture(diag_file, width=Inches(5.5))
        doc.add_paragraph()

    except Exception:
        traceback.print_exc()


# ============================================================
# 11.0 SIMULATION REPORT
# ============================================================

def _add_simulation_report(doc: docx.Document, simulation_results: dict) -> None:
    """
    11.0 Process Performance Report — ISO formatted.
    """
    try:
        if not isinstance(simulation_results, dict) or not simulation_results:
            return

        time_unit = str(simulation_results.get("time_unit", "units"))

        doc.add_heading("11.0 Process Performance Report", level=1)
        doc.add_paragraph(
            f"The following metrics are based on a Monte Carlo discrete-event simulation "
            f"running 2,000 iterations. All time-based values are reported in {time_unit}."
        )

        # Summary table
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Operational Metric"
        hdr[1].text = "Simulated Value"

        metrics_to_show = [
            ("avg_cycle_time", "Average Cycle Time"),
            ("cycle_time_variance", "Cycle Time Variance"),
            ("resource_contention_risk", "Contention Risk Profile"),
        ]

        for key, label in metrics_to_show:
            if key in simulation_results:
                row = table.add_row().cells
                row[0].text = label
                val = simulation_results[key]

                if "cycle_time" in key:
                    try:
                        row[1].text = f"{float(val):.2f} {time_unit}"
                    except:
                        row[1].text = f"{val} {time_unit}"
                else:
                    row[1].text = str(val)

        apply_iso_table_formatting(table)
        doc.add_paragraph()

        # Bottlenecks
        if "bottlenecks" in simulation_results:
            doc.add_heading("Identified Bottlenecks", level=2)
            for b in simulation_results["bottlenecks"]:
                doc.add_paragraph(str(b), style="List Bullet")

        # Per-step performance
        if "per_step_avg" in simulation_results:
            doc.add_heading("Detailed Step Performance", level=2)

            table2 = doc.add_table(rows=1, cols=2)
            hdr2 = table2.rows[0].cells
            hdr2[0].text = "Process Step"
            hdr2[1].text = f"Avg. Duration ({time_unit})"

            for step, avg in simulation_results["per_step_avg"].items():
                row = table2.add_row().cells
                row[0].text = str(step)
                try:
                    row[1].text = f"{float(avg):.2f}"
                except:
                    row[1].text = str(avg)

            apply_iso_table_formatting(table2)
            doc.add_paragraph()

        # Recommendations
        if "recommendations" in simulation_results and isinstance(simulation_results["recommendations"], list):
            doc.add_heading("Optimization Recommendations", level=2)
            for rec in simulation_results["recommendations"]:
                if isinstance(rec, dict):
                    step_name = rec.get("step_name", "Step")
                    instruction = rec.get("instruction", "")
                    doc.add_paragraph(f"{step_name}: {instruction}", style="List Bullet")

    except Exception as e:
        logger.error(f"Error rendering simulation report: {e}")
        traceback.print_exc()

def _add_tools_section_from_summary(doc: docx.Document, tools_summary) -> None:
    """4.0 Tools section: 'tools_summary' list of {category, tools}."""
    try:
        if not isinstance(tools_summary, list) or not tools_summary:
            return

        doc.add_heading("4.0 Supporting Systems and Tools", level=1)
        doc.add_paragraph("The following tools and platforms support this process:")

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Category"
        hdr_cells[1].text = "Tools"

        for entry in tools_summary:
            category = entry.get("category", "")
            tools = entry.get("tools", [])

            row_cells = table.add_row().cells
            row_cells[0].text = str(category)

            if isinstance(tools, list):
                row_cells[1].text = ", ".join(str(x) for x in tools)
            else:
                row_cells[1].text = str(tools)

        apply_iso_table_formatting(table)
        doc.add_paragraph()

    except Exception:
        traceback.print_exc()
