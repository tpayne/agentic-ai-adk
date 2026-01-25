# process_agents/helpers/doc_structure.py

import docx
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import traceback
import logging

logger = logging.getLogger("ProcessArchitect.DocStructure")

def _add_header(doc, label):
    """Adds a bold section sub-header with standard spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.add_run(label).bold = True
    return p


def _add_bullet(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(f"• {text}")

def apply_iso_table_formatting(table: docx.table.Table) -> None:
    """
    Apply a consistent ISO-style formatting to a table:
    - Calibri body style via Normal
    - Light grey header shading
    - Thin borders (via Table Grid style)
    """
    try:
        # Ensure table style is grid-based
        table.style = "Table Grid"

        # Header row shading (10% grey)
        if table.rows:
            hdr_cells = table.rows[0].cells
            for cell in hdr_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9D9D9")  # light grey
                tcPr.append(shd)

        # Ensure all paragraphs use Normal style for font consistency
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.style is None or p.style.name == "Normal":
                        p.style = table._tbl.doc.styles["Normal"]
    except Exception:
        traceback.print_exc()


def add_iso_page_break(doc: docx.Document) -> None:
    """
    Insert a page break with controlled spacing so we don't accumulate
    random blank lines between sections.
    """
    try:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        doc.add_page_break()
    except Exception:
        traceback.print_exc()


def _add_version_history_table(doc: docx.Document, version: str, author: str) -> None:
    """Add a basic version history table derived from JSON or defaults."""
    try:
        doc.add_heading("Document Control", level=1)

        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Version"
        hdr_cells[1].text = "Date"
        hdr_cells[2].text = "Author"
        hdr_cells[3].text = "Description"

        row_cells = table.add_row().cells
        row_cells[0].text = str(version)
        row_cells[1].text = datetime.now().strftime("%Y-%m-%d")
        row_cells[2].text = str(author)
        row_cells[3].text = "Initial generated process specification"

        apply_iso_table_formatting(table)
        doc.add_paragraph()  # spacer
    except Exception:
        traceback.print_exc()


def _add_table_of_contents(doc: docx.Document) -> None:
    """
    Insert a Word Table of Contents field (updates inside Word).
    The user must right-click and 'Update Field' after opening.
    """
    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        fld_simple = OxmlElement("w:fldSimple")
        fld_simple.set(
            qn("w:instr"),
            'TOC \\o "1-3" \\h \\z \\u',
        )
        run._r.append(fld_simple)

        inner_run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "Right-click and select 'Update Field' to generate the Table of Contents."
        inner_run.append(t)
        fld_simple.append(inner_run)
    except Exception:
        traceback.print_exc()
