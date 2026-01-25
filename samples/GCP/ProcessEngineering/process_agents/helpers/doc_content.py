# process_agents/helpers/doc_content.py

import docx
import os
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import traceback
import logging

from .doc_structure import (
    _add_header,
    _add_bullet,
    apply_iso_table_formatting,
    add_iso_page_break,
)

from ..step_diagram_agent import generate_step_diagram_for_step

logger = logging.getLogger("ProcessArchitect.DocContent")


# ============================================================
# 1.0 PROCESS OVERVIEW
# ============================================================

def _add_overview_section(doc: docx.Document, data: dict) -> None:
    """
    1.0 Process Overview — ISO formatted.
    """
    try:
        doc.add_heading("1.0 Process Overview", level=1)

        introduction = data.get("introduction")
        description = data.get("description") or data.get("process_description")

        if introduction:
            doc.add_paragraph(str(introduction))
        elif description:
            doc.add_paragraph(str(description))
        else:
            doc.add_paragraph("This section provides a high-level overview of the business process.")

        subsection = 1

        # --- Assumptions ---
        assumptions = data.get("assumptions")
        if isinstance(assumptions, list) and assumptions:
            doc.add_heading(f"1.{subsection} Assumptions", level=2)
            subsection += 1
            for item in assumptions:
                doc.add_paragraph(item, style="List Bullet")

        # --- Constraints ---
        constraints = data.get("constraints")
        if isinstance(constraints, list) and constraints:
            doc.add_heading(f"1.{subsection} Constraints", level=2)
            subsection += 1
            for item in constraints:
                doc.add_paragraph(item, style="List Bullet")

        # --- Purpose, Scope, Industry ---
        ordered = [
            ("purpose", "Purpose"),
            ("scope", "Scope"),
            ("industry_sector", "Industry Sector"),
        ]

        for key, label in ordered:
            value = data.get(key)
            if value:
                doc.add_heading(f"1.{subsection} {label}", level=2)
                subsection += 1
                doc.add_paragraph(str(value))

        # --- Additional metadata ---
        for key in ["out_of_scope", "business_unit", "owner"]:
            if key in data:
                p = doc.add_paragraph()
                r = p.add_run(f"{key.replace('_', ' ').title()}: ")
                r.bold = True
                p.add_run(str(data.get(key)))

    except Exception:
        traceback.print_exc()


# ============================================================
# 2.0 STAKEHOLDERS
# ============================================================

def _add_stakeholders_section(doc: docx.Document, stakeholders) -> None:
    """
    2.0 Stakeholders — ISO formatted.
    """
    try:
        if not stakeholders or not isinstance(stakeholders, list):
            return

        doc.add_heading("2.0 Stakeholders and Responsibilities", level=1)
        doc.add_paragraph(
            "The following is a list of key stakeholders involved in this process. "
            "Understanding their roles and responsibilities is crucial for successful implementation."
        )

        # Simple list
        if all(isinstance(s, str) for s in stakeholders):
            for s in stakeholders:
                doc.add_paragraph(str(s), style="List Bullet")
            doc.add_paragraph()
            return

        # Table
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Stakeholder"
        hdr[1].text = "Responsibilities"

        for s in stakeholders:
            if not isinstance(s, dict):
                continue
            name = (
                s.get("stakeholder_name")
                or s.get("role_name")
                or s.get("name")
                or s.get("role")
                or "Stakeholder"
            )
            responsibilities = s.get("responsibilities", [])

            row = table.add_row().cells
            row[0].text = str(name)
            if isinstance(responsibilities, list):
                row[1].text = "\n".join(str(x) for x in responsibilities)
            else:
                row[1].text = str(responsibilities)

        apply_iso_table_formatting(table)
        doc.add_paragraph()

    except Exception:
        traceback.print_exc()


# ============================================================
# 3.0 PROCESS WORKFLOW
# ============================================================
def _add_step_diagram_if_available(
    doc: docx.Document,
    step_name: str,
    subprocess_json: dict,
) -> None:
    """
    Generate and embed a subprocess diagram for the given step, if possible.
    Uses the micro-BPMN generator by default.
    """
    try:
        diagram_path = generate_step_diagram_for_step(step_name, subprocess_json)
        if not diagram_path:
            return
        if not os.path.exists(diagram_path):
            return

        doc.add_picture(diagram_path, width=Inches(5.5))
        doc.add_paragraph()  # spacer
    except Exception:
        traceback.print_exc()

def _add_process_steps_section(doc: docx.Document, steps) -> None:
    """
    Hybrid prose + bullet renderer for top-level process steps (3.x).
    No tables. No HTML. Deterministic formatting.
    """
    logger.debug("Rendering process workflow (prose + bullets)…")

    if not isinstance(steps, list) or not steps:
        return

    doc.add_heading("3.0 Process Workflow", level=1)
    doc.add_paragraph(
        "The following is a list of key steps in the process workflow."
    )

    INTRO = {
        "inputs": "The following inputs are required for this step:",
        "outputs": "This step produces the following outputs:",
        "success_criteria": "Success for this step is measured by:",
        "process_triggers": "This step is initiated by:",
        "process_end_conditions": "This step is considered complete when:",
        "dependencies": "This step depends on the following:",
        "deliverables": "This step produces the following deliverables:",
        "governance_requirements": "The following governance requirements apply:",
        "risks_and_controls": "The following risks and controls apply:",
        "step_risks_and_controls": "The following risks and controls apply:",
        "change_management": "The following change management rules apply:",
        "continuous_improvement": "The following continuous improvement practices apply:",
        "estimated_duration": "The estimated duration for this step is:",
        "process_owner": "The following process owner is accountable:",
        "responsible_party": "The following parties are responsible for this step:",
    }

    def expand_value(doc, value, indent=False):
        if isinstance(value, dict):
            for k, v in value.items():
                if isinstance(v, list):
                    _add_bullet(doc, f"{k.replace('_',' ').title()}:", indent)
                    for item in v:
                        _add_bullet(doc, item, indent=True)
                else:
                    _add_bullet(doc, f"{k.replace('_',' ').title()}: {v}", indent)

        elif isinstance(value, list):
            for item in value:
                expand_value(doc, item, indent)

        else:
            _add_bullet(doc, value, indent)

    for s_idx, step in enumerate(steps, start=1):
        if not isinstance(step, dict):
            continue

        step_name = step.get("step_name", f"Step {s_idx}")

        if s_idx > 1:
            add_iso_page_break(doc)

        doc.add_heading(f"3.{s_idx} {step_name}", level=2)

        def prose(label, value):
            if not value:
                return
            doc.add_heading(f"{label}:", level=4)
            doc.add_paragraph(str(value))

        prose("Description", step.get("description"))
        prose("Purpose", step.get("purpose"))
        prose("Scope", step.get("scope"))

        def bullets(field_label, json_key):
            value = step.get(json_key) or step.get(f"step_{json_key}")
            if not value:
                return

            doc.add_heading(f"{field_label}:", level=4)
            doc.add_paragraph(INTRO[json_key])
            expand_value(doc, value)

        bullets("Inputs", "inputs")
        bullets("Outputs", "outputs")
        bullets("Success Criteria", "success_criteria")
        bullets("Process Triggers", "process_triggers")
        bullets("Process End Conditions", "process_end_conditions")
        bullets("Dependencies", "dependencies")
        bullets("Deliverables", "deliverables")
        bullets("Governance Requirements", "governance_requirements")
        bullets("Risks and Controls", "risks_and_controls")
        bullets("Change Management", "change_management")
        bullets("Continuous Improvement", "continuous_improvement")
        bullets("Estimated Duration", "estimated_duration")
        bullets("Process Owner", "process_owner")
        bullets("Responsible Parties", "responsible_party")

        subprocess_json = step.get("subprocess")
        if isinstance(subprocess_json, dict):
            _add_subprocess_section(doc, s_idx, step_name, subprocess_json)

        doc.add_paragraph()

# ============================================================
# SUBPROCESS RENDERING (3.x.y.z)
# ============================================================

def _add_subprocess_section(doc, step_index: int, step_name: str, subprocess_json: dict) -> None:
    """
    Hybrid prose + bullet renderer for subprocess steps (3.x.y).
    Restores diagrams. No tables. No HTML.
    """
    flow = subprocess_json.get("subprocess_flow")
    if not isinstance(flow, list) or not flow:
        return

    add_iso_page_break(doc)

    doc.add_heading(
        f'Required Sub Process(es) for the Step "{step_name}"',
        level=3,
    )
    doc.add_paragraph(
        f'The following details the subprocess flows for the step "{step_name}".'
    )

    _add_step_diagram_if_available(doc, step_name, subprocess_json)

    INTRO = {
        "inputs": "The following inputs are required for this subprocess:",
        "outputs": "This subprocess produces the following outputs:",
        "success_criteria": "Success for this subprocess is measured by:",
        "triggers": "This subprocess is initiated by:",
        "end_conditions": "This subprocess is considered complete when:",
        "dependencies": "This subprocess depends on the following:",
        "governance_requirements": "The following governance requirements apply:",
        "risks_and_controls": "The following risks and controls apply:",
        "step_risks_and_controls": "The following risks and controls apply:",
        "change_management": "The following change management rules apply:",
        "continuous_improvement": "The following continuous improvement practices apply:",
        "estimated_duration": "The estimated duration for this subprocess is:",
        "process_owner": "The following process owner is accountable:",
        "responsible_party": "The following parties are responsible for this subprocess:",
    }

    def expand_value(doc, value, indent=False):
        if isinstance(value, dict):
            for k, v in value.items():
                if isinstance(v, list):
                    _add_bullet(doc, f"{k.replace('_',' ').title()}:", indent)
                    for item in v:
                        _add_bullet(doc, item, indent=True)
                else:
                    _add_bullet(doc, f"{k.replace('_',' ').title()}: {v}", indent)

        elif isinstance(value, list):
            for item in value:
                expand_value(doc, item, indent)

        else:
            _add_bullet(doc, value, indent)

    for sub_idx, sub in enumerate(flow, start=1):
        if not isinstance(sub, dict):
            continue

        sub_name = sub.get("substep_name", f"Sub-step {sub_idx}")

        add_iso_page_break(doc)
        doc.add_heading(f"3.{step_index}.{sub_idx} {sub_name}", level=4)

        doc.add_paragraph(
            f"This subprocess describes the activities required to complete '{sub_name}'."
        )

        diagram = sub.get("diagram")
        if diagram and os.path.exists(diagram):
            doc.add_picture(diagram, width=Inches(6))
            doc.add_paragraph()

        def prose(label, value):
            if not value:
                return
            doc.add_heading(f"{label}:", level=5)
            doc.add_paragraph(str(value))

        prose("Description", sub.get("description"))
        prose("Purpose", sub.get("purpose"))
        prose("Scope", sub.get("scope"))

        def bullets(field_label, json_key):
            value = sub.get(json_key) or sub.get(f"step_{json_key}")
            if not value:
                return

            doc.add_heading(f"{field_label}:", level=5)
            doc.add_paragraph(INTRO[json_key])
            expand_value(doc, value)

        bullets("Inputs", "inputs")
        bullets("Outputs", "outputs")
        bullets("Success Criteria", "success_criteria")
        bullets("Triggers", "triggers")
        bullets("End Conditions", "end_conditions")
        bullets("Dependencies", "dependencies")
        bullets("Governance Requirements", "governance_requirements")
        bullets("Risks and Controls", "risks_and_controls")
        bullets("Change Management", "change_management")
        bullets("Continuous Improvement", "continuous_improvement")
        bullets("Estimated Duration", "estimated_duration")
        bullets("Process Owner", "process_owner")
        bullets("Responsible Party", "responsible_party")

        doc.add_paragraph()


def _render_generic_value(doc: docx.Document, value, label=None) -> None:
    """
    Deterministic renderer: always produces real Word tables for lists/dicts.
    Never prints raw HTML. Never prints raw JSON.
    Mirrors the Stakeholder table logic.
    """

    # ---------------------------
    # Simple string → paragraph
    # ---------------------------
    if isinstance(value, str):
        p = doc.add_paragraph()
        if label:
            r = p.add_run(f"{label}: ")
            r.bold = True
        p.add_run(value)
        return

    # ---------------------------
    # List of simple values
    # ---------------------------
    if isinstance(value, list) and all(isinstance(x, (str, int, float)) for x in value):
        if label:
            doc.add_heading(label, level=3)
        for item in value:
            doc.add_paragraph(str(item), style="List Bullet")
        return

    # ---------------------------
    # List of dicts → table
    # ---------------------------
    if isinstance(value, list) and all(isinstance(x, dict) for x in value):
        if label:
            doc.add_heading(label, level=3)

        # Collect all keys
        all_keys = set()
        for item in value:
            all_keys.update(item.keys())

        ordered_keys = sorted(all_keys)

        table = doc.add_table(rows=1, cols=len(ordered_keys))
        hdr = table.rows[0].cells
        for i, key in enumerate(ordered_keys):
            hdr[i].text = key.replace("_", " ").title()

        for item in value:
            row = table.add_row().cells
            for i, key in enumerate(ordered_keys):
                v = item.get(key, "")
                if isinstance(v, list):
                    row[i].text = "\n".join(str(x) for x in v)
                else:
                    row[i].text = str(v)

        apply_iso_table_formatting(table)
        doc.add_paragraph()
        return

    # ---------------------------
    # Dict → 2‑column table
    # ---------------------------
    if isinstance(value, dict):
        if label:
            doc.add_heading(label, level=3)

        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"

        for k, v in value.items():
            row = table.add_row().cells
            row[0].text = k.replace("_", " ").title()
            if isinstance(v, list):
                row[1].text = "\n".join(str(x) for x in v)
            else:
                row[1].text = str(v)

        apply_iso_table_formatting(table)
        doc.add_paragraph()
        return

    # ---------------------------
    # Fallback
    # ---------------------------
    doc.add_paragraph(str(value))
