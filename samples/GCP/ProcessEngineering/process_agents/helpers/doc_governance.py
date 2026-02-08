# process_agents/helpers/doc_governance.py

import docx
import traceback
import logging

from .doc_structure import (
    _add_header,
    _add_bullet,
    apply_iso_table_formatting,
)
from .doc_content import _render_generic_value

logger = logging.getLogger("ProcessArchitect.DocGovernance")


# ============================================================
# 12.0 GOVERNANCE REQUIREMENTS
# ============================================================

def _add_governance_requirements_section(doc, items):
    """12.0 Governance Requirements — ISO formatted."""
    doc.add_heading("12.0 Governance Requirements", level=1)

    if not items:
        doc.add_paragraph("There are no governance requirements to document.")
        return

    doc.add_paragraph("The following governance requirements apply to this process:")

    for item in items:
        _add_bullet(doc, item)


# ============================================================
# 13.0 RISKS AND CONTROLS
# ============================================================

def _add_risks_and_controls_section(doc, items):
    """13.0 Risks & Controls — ISO formatted table."""
    doc.add_heading("13.0 Risks and Controls", level=1)
    doc.add_paragraph("The following risks and associated controls apply to this process:")

    if not items:
        return

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Risk"
    hdr[1].text = "Control"

    for rc in items:
        row = table.add_row().cells
        row[0].text = str(rc.get("risk", ""))
        row[1].text = str(rc.get("control", ""))

    apply_iso_table_formatting(table, doc)
    doc.add_paragraph()


# ============================================================
# 14.0 PROCESS TRIGGERS
# ============================================================

def _add_process_triggers_section(doc, items):
    """14.0 Process Triggers — ISO formatted."""
    doc.add_heading("14.0 Process Triggers", level=1)

    if not items:
        doc.add_paragraph("There are no process triggers to document.")
        return

    doc.add_paragraph("The following triggers initiate this process:")

    for item in items:
        _add_bullet(doc, item)


# ============================================================
# 15.0 PROCESS END CONDITIONS
# ============================================================

def _add_process_end_conditions_section(doc, items):
    """15.0 Process End Conditions — ISO formatted."""
    doc.add_heading("15.0 Process End Conditions", level=1)

    if not items:
        doc.add_paragraph("There are no process end conditions to document.")
        return

    doc.add_paragraph("The following conditions indicate completion of the process:")

    for item in items:
        _add_bullet(doc, item)


# ============================================================
# 16.0 CHANGE MANAGEMENT
# ============================================================

def _add_change_management_section(doc, items):
    """16.0 Change Management — ISO formatted."""
    doc.add_heading("16.0 Change Management", level=1)

    if not items:
        doc.add_paragraph("There are no change management items to document.")
        return

    doc.add_paragraph("The following change management practices apply to this process:")

    for cm in items:
        if isinstance(cm, dict):
            crp = cm.get("change_request_process")
            vr = cm.get("versioning_rules")

            if crp:
                _add_bullet(doc, f"Change Request Process: {crp}")
            if vr:
                _add_bullet(doc, f"Versioning Rules: {vr}")


# ============================================================
# 17.0 CONTINUOUS IMPROVEMENT
# ============================================================

def _add_continuous_improvement_section(doc, items):
    """17.0 Continuous Improvement — ISO formatted."""
    doc.add_heading("17.0 Continuous Improvement", level=1)

    if not items:
        doc.add_paragraph("There are no continuous improvement items to document.")
        return

    doc.add_paragraph("The following continuous improvement practices apply to this process:")

    for ci in items:
        if not isinstance(ci, dict):
            continue

        freq = ci.get("review_frequency")
        inputs = ci.get("improvement_inputs", [])

        if freq:
            _add_bullet(doc, f"Review Frequency: {freq}")

        if inputs:
            _add_header(doc, "Improvement Inputs:")
            for inp in inputs:
                _add_bullet(doc, inp)


# ============================================================
# APPENDIX A — STRUCTURED APPENDIX
# ============================================================

def _add_appendix_from_json(doc: docx.Document, appendix: dict) -> None:
    """Appendix A: structured appendix content — ISO formatted."""
    try:
        if not isinstance(appendix, dict) or not appendix:
            return

        doc.add_heading("Appendix A: Reference Documents", level=1)
        doc.add_paragraph("The following appendix contains reference materials related to the process:")

        for key, val in appendix.items():
            section_title = str(key).replace("_", " ").title()
            doc.add_heading(section_title, level=2)

            if isinstance(val, dict):
                summary = val.get("summary")
                if summary:
                    p = doc.add_paragraph()
                    r = p.add_run("Summary: ")
                    r.bold = True
                    p.add_run(str(summary))

                last_reviewed = val.get("last_reviewed")
                if last_reviewed:
                    p = doc.add_paragraph()
                    r = p.add_run("Last Reviewed: ")
                    r.bold = True
                    p.add_run(str(last_reviewed))

                review_frequency = val.get("review_frequency")
                if review_frequency:
                    p = doc.add_paragraph()
                    r = p.add_run("Review Frequency: ")
                    r.bold = True
                    p.add_run(str(review_frequency))

                extra = {
                    k: v for k, v in val.items()
                    if k not in {"summary", "last_reviewed", "review_frequency"}
                }
                if extra:
                    _render_generic_value(doc, extra)

            else:
                _render_generic_value(doc, val)

    except Exception:
        traceback.print_exc()


# ============================================================
# APPENDIX B — ADDITIONAL SOURCE DATA
# ============================================================

def _add_additional_data_section(doc: docx.Document, data: dict, consumed_keys: set) -> None:
    """Appendix B: Additional JSON data not covered elsewhere."""
    try:
        if not isinstance(data, dict):
            return

        remaining = {k: v for k, v in data.items() if k not in consumed_keys}
        if not remaining:
            return

        doc.add_heading("Appendix B: Additional Source Data", level=1)
        doc.add_paragraph(
            "This appendix contains additional structured data from the normalized source JSON "
            "that is not covered in the main sections."
        )

        _render_generic_value(doc, remaining)

    except Exception:
        traceback.print_exc()


# ============================================================
# APPENDIX C — GLOSSARY
# ============================================================

def _add_glossary(doc: docx.Document) -> None:
    """Appendix C: Glossary — ISO formatted table."""
    try:
        doc.add_heading("Appendix C: Glossary", level=1)
        doc.add_paragraph("This glossary contains definitions of common terms used in the process documentation:")

        terms = {
            "Business Process": "A set of related activities or tasks performed to achieve a specific organizational goal.",
            "KPI": "Key Performance Indicator used to measure the success or health of a process.",
            "Stakeholder": "Any individual or group with an interest in the outcomes of the process.",
            "Continuous Improvement": "Ongoing effort to improve products, services, or processes.",
        }

        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Term"
        hdr[1].text = "Definition"

        for term, definition in terms.items():
            row = table.add_row().cells
            row[0].text = term
            row[1].text = definition

        apply_iso_table_formatting(table, doc)

    except Exception:
        traceback.print_exc()

def _add_critical_success_factors_section(doc, factors):
    """Adds Critical Success Factors as a table."""
    if not factors:
        return
    
    doc.add_heading("6.0 Critical Success Factors (CSF)", level=1)
    doc.add_paragraph("The following is a list of key CSFs associated with this process.")

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Success Factor"
    hdr[1].text = "Description"

    for factor in factors:
        row = table.add_row().cells
        row[0].text = str(factor.get("name", ""))
        row[1].text = str(factor.get("description", ""))

    apply_iso_table_formatting(table, doc)
    doc.add_paragraph()


def _add_critical_failure_factors_section(doc, factors):
    """Adds Critical Failure Factors as a table."""
    if not factors:
        return
    
    doc.add_heading("7.0 Critical Failure Factors (CFF)", level=1)
    doc.add_paragraph("The following is a list of key CFFs associated with this process.")

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Failure Factor"
    hdr[1].text = "Description"

    for factor in factors:
        row = table.add_row().cells
        row[0].text = str(factor.get("name", ""))
        row[1].text = str(factor.get("description", ""))

    apply_iso_table_formatting(table, doc)
    doc.add_paragraph()

def _add_reporting_and_analytics(doc, items):
    """8.0 Reporting and Analytics — ISO formatted."""
    doc.add_heading("8.0 Reporting and Analytics", level=1)

    if not items:
        doc.add_paragraph("There are no reporting or analytics items to document.")
        return

    doc.add_paragraph(
        "The following reporting and analytics capabilities support monitoring and decision-making:"
    )

    # If list of strings
    if all(isinstance(i, str) for i in items):
        for i in items:
            _add_bullet(doc, i)
        doc.add_paragraph()
        return

    # If list of dicts
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Report / Dashboard"
    hdr[1].text = "Description"

    for entry in items:
        if not isinstance(entry, dict):
            continue

        name = entry.get("name") or entry.get("title") or "Report"
        desc = entry.get("description", "")

        row = table.add_row().cells
        row[0].text = str(name)
        row[1].text = str(desc)

    apply_iso_table_formatting(table, doc)
    doc.add_paragraph()
