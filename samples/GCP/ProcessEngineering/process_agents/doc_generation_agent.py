# process_agents/doc_generation_agent.py

from google.adk.agents import Agent
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .utils import load_instruction, getProperty

import os
import json
from datetime import datetime
import traceback
import logging
import random
import time

logger = logging.getLogger("ProcessArchitect.DocGen")

# Helper imports
from .helpers.doc_structure import (
    _add_version_history_table,
    _add_table_of_contents,
    add_iso_page_break,
)
from .helpers.doc_content import (
    _add_overview_section,
    _add_stakeholders_section,
    _add_process_steps_section,
)
from .helpers.doc_technical import (
    _add_tools_section_from_summary,
    _add_metrics_section,
    _add_system_requirements,
    _add_flowchart_section,
    _add_simulation_report,
)
from .helpers.doc_governance import (
    _add_governance_requirements_section,
    _add_risks_and_controls_section,
    _add_process_triggers_section,
    _add_process_end_conditions_section,
    _add_change_management_section,
    _add_continuous_improvement_section,
    _add_appendix_from_json,
    _add_additional_data_section,
    _add_glossary,
)


# ============================================================
# LOAD SUBPROCESSES
# ============================================================

def _load_subprocesses() -> dict:
    """Loads all subprocess JSON files from output/subprocesses/."""
    logger.debug("Loading subprocess JSON files...")
    subprocess_dir = "output/subprocesses"
    subprocesses = {}

    if not os.path.isdir(subprocess_dir):
        return subprocesses

    for filename in os.listdir(subprocess_dir):
        if not filename.endswith(".json"):
            continue

        path = os.path.join(subprocess_dir, filename)
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)

            parent = data.get("parent_step_name") or data.get("step_name")
            if parent:
                data["step_name"] = parent
                subprocesses[parent] = data

        except Exception:
            logger.exception(f"Failed to load subprocess file: {path}")

    logger.debug("Subprocesses loaded.")
    return subprocesses


# ============================================================
# GLOBAL DOCUMENT STYLE SETUP
# ============================================================

def _apply_global_styles(doc: docx.Document):
    """Apply ISO-style global typography and spacing."""
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)

    # Heading 4
    h4 = styles["Heading 4"]
    h4.font.name = "Calibri"
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.paragraph_format.space_before = Pt(6)
    h4.paragraph_format.space_after = Pt(3)

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)


# ============================================================
# MAIN DOCUMENT GENERATION
# ============================================================

def create_standard_doc_from_file(process_name: str) -> str:
    """
    Generate a structured, ISO-formatted process document from process_data.json.
    """
    time.sleep(float(getProperty("modelSleep")) + random.random() * 0.75)
    logger.debug(f"Creating document for process: {process_name}...")

    try:
        with open("output/process_data.json", "r", encoding="utf-8") as f:
            raw_data = json.load(f)

        if isinstance(raw_data, dict) and "process_design" in raw_data:
            data = raw_data["process_design"]
        else:
            data = raw_data

        # Attach subprocesses
        subprocesses = _load_subprocesses()
        for step in data.get("process_steps", []):
            if isinstance(step, dict):
                name = step.get("step_name")
                if name in subprocesses:
                    step["subprocess"] = subprocesses[name]

        # Extract metadata
        name = str(data.get("process_name", process_name))
        description = data.get("description") or data.get("process_description")
        version = str(data.get("version", "1.0"))
        sector = data.get("industry_sector", data.get("business_unit", "N/A"))

        stakeholders = data.get("stakeholders")
        process_steps = data.get("process_steps")
        tools_summary = data.get("tools_summary")
        critical_success_factors = data.get("critical_success_factors")
        critical_failure_factors = data.get("critical_failure_factors")
        metrics = data.get("metrics") or data.get("success_metrics")
        reporting_and_analytics = data.get("reporting_and_analytics")
        system_requirements = data.get("system_requirements")
        appendix = data.get("appendix") if isinstance(data.get("appendix"), dict) else None

        governance_requirements = data.get("governance_requirements")
        process_end_conditions = data.get("process_end_conditions")
        change_management = data.get("change_management")
        process_triggers = data.get("process_triggers")
        continuous_improvement = data.get("continuous_improvement")
        risks_and_controls = data.get("risks_and_controls")

        consumed_keys = {
            "appendix", "assumptions", "business_unit", "change_management",
            "constraints", "continuous_improvement", "critical_failure_factors",
            "critical_success_factors", "description", "governance_requirements",
            "industry_sector", "introduction", "metrics", "process_description",
            "process_end_conditions", "process_name", "process_steps",
            "process_triggers", "reporting_and_analytics", "risks_and_controls",
            "stakeholders", "success_metrics", "system_requirements",
            "tools_summary", "version",
        }

        # Create document
        from process_agents.helpers.themes import apply_theme 
        doc = docx.Document()
        theme = getProperty("theme")
        if theme:
            apply_theme(doc, theme)
        else:
            _apply_global_styles(doc)

        # ============================================================
        # COVER PAGE
        # ============================================================

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(28)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("End-to-End Level 0 Value Stream Map").font.size = Pt(18)

        doc.add_paragraph()
        doc.add_paragraph(f"Industry / Domain: {sector}")
        doc.add_paragraph(f"Version: {version}")
        doc.add_paragraph(f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        add_iso_page_break(doc)

        # ============================================================
        # DOCUMENT CONTROL + TOC
        # ============================================================

        _add_version_history_table(doc, version=version, author="Process Architect")

        add_iso_page_break(doc)
        doc.add_heading("Table of Contents", level=1)
        _add_table_of_contents(doc)

        add_iso_page_break(doc)

        # ============================================================
        # MAIN SECTIONS
        # ============================================================

        _add_overview_section(doc, data)
        add_iso_page_break(doc)

        _add_stakeholders_section(doc, stakeholders)
        add_iso_page_break(doc)

        if process_steps:
            _add_process_steps_section(doc, process_steps)
            add_iso_page_break(doc)

        if tools_summary:
            _add_tools_section_from_summary(doc, tools_summary)
            add_iso_page_break(doc)

        if isinstance(metrics, list) and metrics:
            _add_metrics_section(doc, metrics)
            add_iso_page_break(doc)

        if isinstance(critical_success_factors, list) and critical_success_factors:
            from .helpers.doc_governance import _add_critical_success_factors_section
            _add_critical_success_factors_section(doc, critical_success_factors)
            add_iso_page_break(doc)

        if isinstance(critical_failure_factors, list) and critical_failure_factors:
            from .helpers.doc_governance import _add_critical_failure_factors_section
            _add_critical_failure_factors_section(doc, critical_failure_factors)
            add_iso_page_break(doc)

        if reporting_and_analytics:
            from .helpers.doc_governance import _add_reporting_and_analytics
            _add_reporting_and_analytics(doc, reporting_and_analytics)
            add_iso_page_break(doc)

        if system_requirements:
            _add_system_requirements(doc, system_requirements)
            add_iso_page_break(doc)

        _add_flowchart_section(doc, name)
        add_iso_page_break(doc)

        # Simulation results
        simulation_results = None
        try:
            sim_path = "output/simulation_results.json"
            if os.path.exists(sim_path):
                with open(sim_path, "r", encoding="utf-8") as sf:
                    simulation_results = json.load(sf)
        except Exception:
            traceback.print_exc()

        if simulation_results:
            _add_simulation_report(doc, simulation_results)
            add_iso_page_break(doc)

        if governance_requirements:
            _add_governance_requirements_section(doc, governance_requirements)
            add_iso_page_break(doc)

        if risks_and_controls:
            _add_risks_and_controls_section(doc, risks_and_controls)
            add_iso_page_break(doc)

        if process_triggers:
            _add_process_triggers_section(doc, process_triggers)
            add_iso_page_break(doc)

        if process_end_conditions:
            _add_process_end_conditions_section(doc, process_end_conditions)
            add_iso_page_break(doc)

        if change_management:
            _add_change_management_section(doc, change_management)
            add_iso_page_break(doc)

        if continuous_improvement:
            _add_continuous_improvement_section(doc, continuous_improvement)
            add_iso_page_break(doc)

        # Appendices
        if appendix:
            _add_appendix_from_json(doc, appendix)
            consumed_keys.add("appendix")

        add_iso_page_break(doc)
        _add_additional_data_section(doc, data, consumed_keys)

        add_iso_page_break(doc)
        _add_glossary(doc)

        # Save
        out_path = f"output/{name.replace(' ', '_')}.docx"
        doc.save(out_path)
        return f"SUCCESS: Professional document saved at {out_path}"

    except Exception as e:
        traceback.print_exc()
        return f"ERROR: {str(e)}"


# ============================================================
# AGENT DEFINITION
# ============================================================

doc_generation_agent = Agent(
    name="Document_Generation_Agent",
    description="Generates a professional ISO-formatted Word document from normalized JSON.",
    instruction=load_instruction("doc_generation_agent.txt"),
    tools=[create_standard_doc_from_file],
)

if __name__ == "__main__":
    import sys

    # Allow: python -m process_agents.doc_generation_agent <path>
    if len(sys.argv) > 1:
        input_path = sys.argv[1]
    else:
        input_path = "output/process_data.json"

    try:
        result = create_standard_doc_from_file(input_path)
        print(result)
    except Exception as e:
        print(f"ERROR during document generation: {e}")
        raise
